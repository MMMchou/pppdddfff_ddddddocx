#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查 Word 文档的页数和内容
"""

from docx import Document
from pathlib import Path
import sys


def count_docx_content(docx_path):
    """统计 Word 文档的内容"""
    doc = Document(docx_path)
    
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    
    # 统计文本内容
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    text_length = len(text_content.strip())
    
    # 统计分页符
    page_breaks = sum(1 for para in doc.paragraphs if para.text == '' and hasattr(para, '_element') and 'page-break' in str(para._element.xml))
    
    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'text_length': text_length,
        'page_breaks': page_breaks
    }


def compare_docx_files(docx_dir):
    """对比目录中的多个 Word 文档"""
    dir_path = Path(docx_dir)
    
    if not dir_path.exists():
        print(f"目录不存在: {docx_dir}")
        return
    
    docx_files = sorted(dir_path.glob("*.docx"))
    
    if not docx_files:
        print(f"未找到 Word 文档")
        return
    
    print("=" * 80)
    print(f"📊 Word 文档对比: {dir_path.name}")
    print("=" * 80)
    print()
    
    results = []
    for docx_file in docx_files:
        try:
            stats = count_docx_content(docx_file)
            file_size = docx_file.stat().st_size / 1024
            
            results.append({
                'name': docx_file.name,
                'size': file_size,
                **stats
            })
        except Exception as e:
            print(f"⚠️  {docx_file.name}: 无法读取 - {e}")
    
    if not results:
        return
    
    # 打印表格
    print(f"{'文件名':<40} {'大小(KB)':<10} {'段落':<8} {'表格':<8} {'文本长度':<10}")
    print("-" * 80)
    
    for r in results:
        print(f"{r['name']:<40} {r['size']:>8.1f}  {r['paragraphs']:>6}   {r['tables']:>6}   {r['text_length']:>8}")
    
    print()
    print("=" * 80)
    
    # 提示
    print("\n💡 提示:")
    print("  - 文本长度相同 = 内容完整一致")
    print("  - 段落数相近 = 格式基本保持")
    print("  - 推荐使用: *_simple.docx（简单拼接，无多余页）")
    print()


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python check_pages.py <docx目录>")
        print("示例: python check_pages.py output/常规2/final")
        sys.exit(1)
    
    compare_docx_files(sys.argv[1])

