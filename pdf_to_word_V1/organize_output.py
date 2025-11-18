#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
整理和合并 PaddleOCR 输出结果
自动将分散的页面合并成最终文档
"""

import os
import shutil
from pathlib import Path
from docx import Document  # pyright: ignore[reportMissingImports]
from docx.opc.exceptions import PackageNotFoundError  # pyright: ignore[reportMissingImports]
import json


def merge_docx_files(docx_files, output_path):
    """
    合并多个 Word 文档，使用 docxcompose 保持样式一致性
    
    特点：
    - 使用分节符（Section Break）而不是简单的分页符
    - 保留所有格式（字体、颜色、表格样式、图片等）
    - 自动处理样式冲突
    - 每页独立成一节
    """
    if not docx_files:
        return False
    
    try:
        from docxcompose.composer import Composer  # pyright: ignore[reportMissingImports]
        
        # 创建主文档（第一页）
        master = Document(docx_files[0])
        composer = Composer(master)
        
        # 依次追加其他文档
        for docx_file in docx_files[1:]:
            try:
                # 在追加新文档前，先添加分页符
                master.add_page_break()
                
                # 使用 composer 追加文档（保持样式一致性）
                doc = Document(docx_file)
                composer.append(doc)
            except Exception as e:
                print(f"警告: 无法合并 {docx_file}: {e}")
        
        # 保存合并后的文档
        composer.save(output_path)
        return True
        
    except ImportError:
        # 如果 docxcompose 未安装，回退到基础方法
        print("提示: docxcompose 未安装，使用基础合并方法")
        return _merge_docx_basic(docx_files, output_path)
    except Exception as e:
        print(f"docxcompose 合并失败: {e}，尝试基础方法")
        return _merge_docx_basic(docx_files, output_path)


def _merge_docx_basic(docx_files, output_path):
    """
    基础合并方法（备用）
    使用 XML 元素复制
    """
    if not docx_files:
        return False
    
    from copy import deepcopy
    
    # 创建主文档
    merged_doc = Document(docx_files[0])
    
    # 添加其他文档
    for docx_file in docx_files[1:]:
        try:
            doc = Document(docx_file)
            
            # 添加分页符
            merged_doc.add_page_break()
            
            # 复制 XML 元素
            for element in doc.element.body:
                element_copy = deepcopy(element)
                merged_doc.element.body.append(element_copy)
                        
        except Exception as e:
            print(f"警告: 无法合并 {docx_file}: {e}")
    
    merged_doc.save(output_path)
    return True


def merge_markdown_files(md_files, output_path):
    """合并多个 Markdown 文件"""
    if not md_files:
        return False
    
    with open(output_path, 'w', encoding='utf-8') as outfile:
        for i, md_file in enumerate(md_files):
            try:
                with open(md_file, 'r', encoding='utf-8') as infile:
                    content = infile.read()
                    if i > 0:
                        outfile.write(f"\n\n---\n\n# 第 {i+1} 页\n\n")
                    outfile.write(content)
            except Exception as e:
                print(f"警告: 无法读取 {md_file}: {e}")
    
    return True


def organize_output_directory(output_dir):
    """
    整理输出目录结构
    
    目录结构:
    output/文件名/
        ├── final/              # 最终合并文档
        │   ├── 文件名.docx
        │   └── 文件名.md
        ├── pages/              # 分页文档
        │   ├── page_0.docx
        │   ├── page_1.docx
        │   └── ...
        ├── images/             # 可视化图片
        │   └── ...
        └── debug/              # 调试信息
            ├── json/
            └── tex/
    """
    output_path = Path(output_dir)
    
    if not output_path.exists():
        print(f"目录不存在: {output_dir}")
        return False
    
    print(f"正在整理: {output_dir}")
    
    # 创建子目录
    final_dir = output_path / "final"
    pages_dir = output_path / "pages"
    images_dir = output_path / "images"
    debug_dir = output_path / "debug"
    
    final_dir.mkdir(exist_ok=True)
    pages_dir.mkdir(exist_ok=True)
    images_dir.mkdir(exist_ok=True)
    (debug_dir / "json").mkdir(parents=True, exist_ok=True)
    (debug_dir / "tex").mkdir(parents=True, exist_ok=True)
    
    # 获取基础文件名
    base_name = output_path.name
    
    # 收集所有文件
    docx_files = []
    md_files = []
    
    # 扫描并分类文件
    for file in sorted(output_path.glob(f"{base_name}_*.docx")):
        page_num = file.stem.split('_')[-1]
        if page_num.isdigit():
            docx_files.append((int(page_num), file))
    
    for file in sorted(output_path.glob(f"{base_name}_*.md")):
        page_num = file.stem.split('_')[-1]
        if page_num.isdigit():
            md_files.append((int(page_num), file))
    
    # 按页码排序
    docx_files.sort(key=lambda x: x[0])
    md_files.sort(key=lambda x: x[0])
    
    print(f"找到 {len(docx_files)} 个 Word 文档")
    print(f"找到 {len(md_files)} 个 Markdown 文档")
    
    # 合并 Word 文档
    if docx_files:
        print("正在合并 Word 文档...")
        docx_paths = [f[1] for f in docx_files]
        final_docx = final_dir / f"{base_name}.docx"
        
        if merge_docx_files(docx_paths, final_docx):
            print(f"✓ Word 文档已合并: {final_docx}")
            
            # 移动分页文档
            for i, (page_num, file) in enumerate(docx_files):
                target = pages_dir / f"page_{page_num}.docx"
                shutil.copy2(file, target)
                file.unlink()  # 删除原文件
    
    # 合并 Markdown 文档
    if md_files:
        print("正在合并 Markdown 文档...")
        md_paths = [f[1] for f in md_files]
        final_md = final_dir / f"{base_name}.md"
        
        if merge_markdown_files(md_paths, final_md):
            print(f"✓ Markdown 文档已合并: {final_md}")
            
            # 移动分页文档
            for i, (page_num, file) in enumerate(md_files):
                target = pages_dir / f"page_{page_num}.md"
                shutil.copy2(file, target)
                file.unlink()  # 删除原文件
    
    # 整理图片文件
    for img_file in output_path.glob("*.png"):
        if img_file.is_file():
            shutil.move(str(img_file), str(images_dir / img_file.name))
    
    # 整理 JSON 文件
    for json_file in output_path.glob("*.json"):
        if json_file.is_file():
            shutil.move(str(json_file), str(debug_dir / "json" / json_file.name))
    
    # 整理 TEX 文件
    for tex_file in output_path.glob("*.tex"):
        if tex_file.is_file():
            shutil.move(str(tex_file), str(debug_dir / "tex" / tex_file.name))
    
    # 移动 imgs 目录
    imgs_src = output_path / "imgs"
    if imgs_src.exists():
        imgs_target = images_dir / "extracted"
        if imgs_target.exists():
            shutil.rmtree(imgs_target)
        shutil.move(str(imgs_src), str(imgs_target))
    
    # 创建 README
    readme_path = final_dir / "README.txt"
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(f"""
==============================================
  {base_name} - 转换结果
==============================================

📁 文件结构:

final/
  ├── {base_name}.docx    ← 最终合并的 Word 文档
  └── {base_name}.md      ← 最终合并的 Markdown 文档

pages/
  ├── page_0.docx         ← 第 1 页（独立）
  ├── page_1.docx         ← 第 2 页（独立）
  └── ...

images/
  ├── *_layout_det_res.png      ← 版面检测结果
  ├── *_overall_ocr_res.png     ← OCR 识别结果
  └── extracted/                ← 提取的图片

debug/
  ├── json/                     ← JSON 数据
  └── tex/                      ← LaTeX 公式

==============================================

✨ 推荐使用:
   - 最终文档: {base_name}.docx
   - Markdown: {base_name}.md

⚠️ 如需查看分页结果，请查看 pages/ 目录

生成时间: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
==============================================
""")
    
    print(f"\n✓ 整理完成！")
    print(f"  最终文档: {final_dir}")
    print(f"  - Word: {final_dir / f'{base_name}.docx'}")
    print(f"  - Markdown: {final_dir / f'{base_name}.md'}")
    
    return True


def organize_all_outputs(base_dir="output"):
    """批量整理指定目录下的所有输出"""
    base_path = Path(base_dir)
    
    if not base_path.exists():
        print(f"目录不存在: {base_dir}")
        return
    
    # 找到所有需要整理的目录
    dirs_to_organize = []
    
    for item in base_path.iterdir():
        if item.is_dir():
            # 检查是否有分页的 docx 文件
            docx_files = list(item.glob("*_*.docx"))
            final_dir = item / "final"
            
            # 如果有分页文件，且 final 目录不存在或为空
            if docx_files and (not final_dir.exists() or not list(final_dir.glob("*.docx"))):
                dirs_to_organize.append(item)
    
    if not dirs_to_organize:
        print("✓ 所有输出目录已整理完成，无需处理")
        return
    
    print(f"找到 {len(dirs_to_organize)} 个目录需要整理:\n")
    for d in dirs_to_organize:
        print(f"  - {d.name}")
    
    print(f"\n开始整理...\n")
    print("=" * 60)
    
    success_count = 0
    for i, output_dir in enumerate(dirs_to_organize, 1):
        print(f"\n[{i}/{len(dirs_to_organize)}] {output_dir.name}")
        print("-" * 60)
        
        try:
            organize_output_directory(str(output_dir))
            success_count += 1
        except Exception as e:
            print(f"✗ 整理失败: {e}")
        
        print("-" * 60)
    
    print("=" * 60)
    print(f"\n整理完成！")
    print(f"  成功: {success_count}/{len(dirs_to_organize)}")
    print(f"  失败: {len(dirs_to_organize) - success_count}/{len(dirs_to_organize)}")


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='整理 PaddleOCR 输出结果')
    parser.add_argument('output_dir', nargs='?', default='output',
                        help='输出目录路径（单个目录或包含多个输出的根目录，默认: output）')
    parser.add_argument('--batch', action='store_true',
                        help='批量整理模式：整理指定目录下的所有输出子目录')
    
    args = parser.parse_args()
    
    if args.batch:
        # 批量整理
        organize_all_outputs(args.output_dir)
    else:
        # 单个目录整理
        organize_output_directory(args.output_dir)


if __name__ == '__main__':
    main()

